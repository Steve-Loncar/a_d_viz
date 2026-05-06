from __future__ import annotations

import io
import math
import numpy as np
import pandas as pd
import streamlit as st
import plotly.graph_objects as go

DOCX_IMPORT_ERROR = None
try:
    from docx import Document
    from docx.shared import Pt
except Exception as e:
    Document = None  # handled at runtime
    Pt = None
    DOCX_IMPORT_ERROR = repr(e)

# Muted red → amber → green (dark-mode friendly)
MUTED_RAG = [
    [0.0,  "#7f1d1d"],
    [0.25, "#9a3412"],
    [0.5,  "#a16207"],
    [0.75, "#166534"],
    [1.0,  "#14532d"],
]

HIER_COLS = [
    "Hierarchy - Level 1",
    "Hierarchy - Level 2",
    "Hierarchy - Level 3",
    "Hierarchy - Level 4",
]

import re
import textwrap
from typing import List
import plotly.express as px
from plotly.subplots import make_subplots
import html

from lib.loader import load_workbook_bytes, load_workbook_path
from lib.transforms import derive_hierarchy, safe_num, clean_players, clean_proxies

st.set_page_config(page_title="A&D Market Explorer (v2)", layout="wide")

# AGGRESSIVE dark theme override for Streamlit + Plotly
st.markdown(
    """
    <style>
    /* Text inputs & text areas – dark background always */
    .stTextInput > div > input,
    .stTextArea > div > textarea {
        background-color: #020617 !important;  /* near-black */
        color: #e5e7eb !important;            /* light gray text */
        border: 1px solid #1f2937 !important; /* subtle border */
    }

    /* Hover + focus: keep dark, just lighten border */
    .stTextInput > div > input:hover,
    .stTextInput > div > input:focus,
    .stTextArea > div > textarea:hover,
    .stTextArea > div > textarea:focus {
        background-color: #020617 !important;
        color: #f9fafb !important;
        border-color: #4b5563 !important;     /* slightly brighter outline */
        box-shadow: none !important;          /* remove default glow */
    }

    /* Placeholder text: lighter but still readable on dark */
    .stTextInput > div > input::placeholder,
    .stTextArea > div > textarea::placeholder {
        color: #6b7280 !important;
    }

    /* NUCLEAR OPTION: Force all Plotly elements dark */
    .js-plotly-plot .plotly,
    .js-plotly-plot .plotly .svg-container,
    .modebar,
    .modebar-container {
        background-color: transparent !important;
    }

    /* Plotly range slider (the white box at top of chart) */
    .rangeslider-container,
    .rangeslider-mask-min,
    .rangeslider-mask-max,
    .range-slider,
    .slider,
    .slider-bg {
        background-color: #0f172a !important;
        fill: #0f172a !important;
    }

    /* Plotly legend boxes */
    .legend,
    .scrollbox {
        background-color: rgba(15, 23, 42, 0.9) !important;
        border: 1px solid rgba(148, 163, 184, 0.35) !important;
    }

    /* Any remaining white backgrounds in Plotly */
    .plotly .bg,
    .plotly .gridlayer,
    rect[fill="#ffffff"],
    rect[fill="rgb(255,255,255)"],
    rect[fill="white"] {
        fill: #0f172a !important;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# Set global Plotly template to dark
import plotly.io as pio
pio.templates.default = "plotly_dark"

def _safe_float(v):
    try:
        if v is None:
            return None
        f = float(v)
        if math.isnan(f):
            return None
        return f
    except Exception:
        return None

CUSTOM_HEATMAP_METRICS = {
    "Margin (FY23)": {"kind": "col", "col": "segment_fy23_ebitda_margin_pct", "unit": "%"},
    "Margin (FY24)": {"kind": "col", "col": "segment_fy24_ebitda_margin_pct", "unit": "%"},
    "Margin (FY25)": {"kind": "col", "col": "segment_fy25_ebitda_margin_pct", "unit": "%"},
    "Revenue Growth 22–25": {"kind": "growth_pct", "start": "segment_fy22_revenue_usd_bn", "end": "segment_fy25_revenue_usd_bn", "unit": "%"},
    "EBITDA Growth 22–25": {"kind": "growth_pct", "start": "segment_fy22_ebitda_usd_bn", "end": "segment_fy25_ebitda_usd_bn", "unit": "%"},
}

def _metric_value(row: pd.Series, spec: dict) -> float | None:
    kind = spec.get("kind")
    if kind == "col":
        return _safe_float(row.get(spec.get("col")))
    if kind == "growth_pct":
        s = _safe_float(row.get(spec.get("start")))
        e = _safe_float(row.get(spec.get("end")))
        if s is None or e is None or s == 0:
            return None
        return (e / s - 1.0) * 100.0
    return None

def _hier_cols_present(nodes: pd.DataFrame) -> list[str]:
    return [c for c in HIER_COLS if c in nodes.columns]

def _node_path_from_levels(row: pd.Series, cols: list[str]) -> str:
    parts = []
    for c in cols:
        v = str(row.get(c, "")).strip()
        if v and v.lower() != "nan":
            parts.append(v)
    return " > ".join(parts)

def _add_unique(existing: list[str], new_items: list[str]) -> list[str]:
    # preserve order, dedupe
    merged = list(existing)
    seen = set(existing)
    for x in new_items:
        if x not in seen:
            merged.append(x)
            seen.add(x)
    return merged

def _norm(s: str) -> str:
    return " ".join(str(s or "").strip().lower().split())

def split_path_levels(path: str, max_levels: int = 6) -> list[str]:
    """
    Split taxonomy path into hierarchy levels.
    """
    if not isinstance(path, str):
        return []
    parts = [p.strip() for p in path.split(">")]
    return parts[:max_levels]

def _first_present(df: pd.DataFrame, cols: list[str]) -> str | None:
    for c in cols:
        if c in df.columns:
            return c
    return None

def _with_path_hierarchy_from_df(df: pd.DataFrame) -> tuple[pd.DataFrame, list[str]]:
    """Derive H1..Hn columns from `path` split on '>'."""
    out = df.copy()
    paths = out["path"].fillna("").astype(str)
    parts = paths.apply(lambda s: [p.strip() for p in re.split(r"\s*>\s*", s) if p.strip()])
    max_depth = int(parts.map(len).max()) if len(parts) else 0
    if max_depth <= 0:
        return out, []
    hcols = []
    for i in range(max_depth):
        col = f"H{i+1}"
        hcols.append(col)
        out[col] = parts.map(lambda xs: xs[i] if i < len(xs) else "")
    return out, hcols

def _pct_growth(v0: float, v1: float) -> float:
    """
    Simple % growth from v0 to v1.
    Returns NaN if baseline is missing/zero.
    """
    try:
        if v0 is None or v1 is None:
            return float("nan")
        if isinstance(v0, float) and math.isnan(v0):
            return float("nan")
        if isinstance(v1, float) and math.isnan(v1):
            return float("nan")
        v0 = float(v0)
        v1 = float(v1)
        if v0 == 0.0:
            return float("nan")
        return (v1 / v0 - 1.0) * 100.0
    except Exception:
        return float("nan")

def _metric_col(metric: str, year: int) -> str:
    """
    Map user-facing metric choice to v2 Nodes column.
    metric: 'Revenue' | 'EBITDA' | 'Margin'
    year: 2023/2024/2025
    """
    yy = str(year)[-2:]
    if metric == "Revenue":
        return f"segment_fy{yy}_revenue_usd_bn"
    if metric == "EBITDA":
        return f"segment_fy{yy}_ebitda_usd_bn"
    return f"segment_fy{yy}_ebitda_margin_pct"

def _path_parts(path: str) -> list[str]:
    return [p.strip() for p in re.split(r"\s*>\s*", str(path or "")) if p.strip()]

def _lookup_with_fallback(metric_map: dict[str, float], path: str) -> float:
    """
    Old-app style fallback: if an ancestor node isn't present in Nodes,
    walk up until you find a value.
    """
    parts = _path_parts(path)
    while parts:
        key = " > ".join(parts)
        v = metric_map.get(key, None)
        if v is not None:
            return v
        parts = parts[:-1]
    return float("nan")

def _with_path_hierarchy(nodes: pd.DataFrame) -> tuple[pd.DataFrame, List[str]]:
    """
    Create temporary hierarchy columns from `path` like:
      H1, H2, H3... (split on '>')
    Returns (df_with_cols, colnames).
    """
    df = nodes.copy()
    paths = df["path"].fillna("").astype(str)
    parts = paths.apply(lambda s: [p.strip() for p in re.split(r"\s*>\s*", s) if p.strip()])
    max_depth = int(parts.map(len).max()) if len(parts) else 0
    if max_depth <= 0:
        return df, []
    colnames = []
    for i in range(max_depth):
        col = f"H{i+1}"
        colnames.append(col)
        df[col] = parts.map(lambda xs: xs[i] if i < len(xs) else "")
    return df, colnames

def _node_index_for_path(nodes: pd.DataFrame, path: str) -> int | None:
    """Return the Nodes dataframe index for an exact path match."""
    if not path:
        return None
    if "path" not in nodes.columns:
        return None
    hit = nodes[nodes["path"].fillna("").astype(str).str.strip() == str(path).strip()]
    if hit.empty:
        return None
    return int(hit.index[0])

def _doc_add_heading(doc: "Document", text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)

def _doc_add_bullets(doc: "Document", lines: list[str]) -> None:
    for ln in [l.strip() for l in lines if str(l).strip()]:
        p = doc.add_paragraph(ln, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10) if Pt else None

def _doc_add_para(doc: "Document", text: str) -> None:
    p = doc.add_paragraph(str(text or "").strip())
    for run in p.runs:
        run.font.size = Pt(10) if Pt else None

def _doc_add_table(doc: "Document", df: pd.DataFrame, title: str | None = None) -> None:
    if title:
        _doc_add_heading(doc, title, level=3)
    if df is None or df.empty:
        _doc_add_para(doc, "—")
        return
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for j, c in enumerate(df.columns):
        hdr[j].text = str(c)
    for _, r in df.iterrows():
        row = t.add_row().cells
        for j, c in enumerate(df.columns):
            row[j].text = "" if pd.isna(r[c]) else str(r[c])

def _fmt_num(v, decimals=2):
    try:
        if v is None or (isinstance(v, float) and math.isnan(v)):
            return "—"
        return f"{float(v):,.{decimals}f}"
    except Exception:
        return "—"

def _node_kpis(node: pd.Series, fy: int = 25) -> dict[str, float | None]:
    rev = safe_num(node.get(f"segment_fy{fy:02d}_revenue_usd_bn"))
    ebitda = safe_num(node.get(f"segment_fy{fy:02d}_ebitda_usd_bn"))
    margin = safe_num(node.get(f"segment_fy{fy:02d}_ebitda_margin_pct"))
    return {"rev": rev, "ebitda": ebitda, "margin": margin}

def _growth_pct(a, b) -> float | None:
    a = safe_num(a)
    b = safe_num(b)
    try:
        if a is None or b is None:
            return None
        if isinstance(a, float) and math.isnan(a):
            return None
        if a == 0:
            return None
        return (float(b) / float(a) - 1.0) * 100.0
    except Exception:
        return None

def build_node_report_docx(
    node: pd.Series,
    nodes: pd.DataFrame,
    players_all: pd.DataFrame,
    proxies_all: pd.DataFrame,
    evidence_all: pd.DataFrame,
    evidence_map_all: pd.DataFrame,
) -> bytes:
    """
    v1 Node report: consultant-style, covers what's in the tabs.
    Returns docx bytes.
    """
    if Document is None:
        raise RuntimeError("python-docx is not installed in this environment.")

    doc = Document()

    node_name = str(node.get("display_name", "") or node.get("node_name", "") or "Node").strip()
    node_path = str(node.get("path", "") or "").strip()
    node_id = str(node.get("node_id", "") or "").strip()

    # ----------------
    # Cover / snapshot
    # ----------------
    _doc_add_heading(doc, f"{node_name}", level=1)
    if node_path:
        _doc_add_para(doc, f"Path: {node_path}")
    if node_id:
        _doc_add_para(doc, f"Node ID: {node_id}")

    k25 = _node_kpis(node, 25)
    k24 = _node_kpis(node, 24)
    rev_g_22_25 = _growth_pct(node.get("segment_fy22_revenue_usd_bn"), node.get("segment_fy25_revenue_usd_bn"))
    ebt_g_22_25 = _growth_pct(node.get("segment_fy22_ebitda_usd_bn"), node.get("segment_fy25_ebitda_usd_bn"))

    _doc_add_heading(doc, "Snapshot (FY2025 default)", level=2)
    _doc_add_bullets(
        doc,
        [
            f"Revenue (USD bn): {_fmt_num(k25['rev'], 3)}",
            f"EBITDA (USD bn): {_fmt_num(k25['ebitda'], 3)}",
            f"EBITDA margin (%): {_fmt_num(k25['margin'], 1)}",
            f"Revenue growth FY22–FY25 (%): {_fmt_num(rev_g_22_25, 1)}",
            f"EBITDA growth FY22–FY25 (%): {_fmt_num(ebt_g_22_25, 1)}",
            f"YoY revenue change FY24→FY25 (%): {_fmt_num(_growth_pct(k24['rev'], k25['rev']), 1)}",
            f"YoY EBITDA change FY24→FY25 (%): {_fmt_num(_growth_pct(k24['ebitda'], k25['ebitda']), 1)}",
        ],
    )

    # ----------------
    # Overview content
    # ----------------
    desc = str(node.get("node_description", "") or node.get("description", "") or "").strip()
    scope = str(node.get("scope_context", "") or "").strip()
    method = str(node.get("methodology_summary", "") or "").strip()

    _doc_add_heading(doc, "Overview", level=2)
    if desc:
        _doc_add_heading(doc, "Node description", level=3)
        _doc_add_para(doc, desc)
    if scope:
        _doc_add_heading(doc, "Scope", level=3)
        _doc_add_para(doc, scope)
    if method:
        _doc_add_heading(doc, "Methodology", level=3)
        _doc_add_para(doc, method)

    # ----------------
    # Node financials (FY22–FY25)
    # ----------------
    _doc_add_heading(doc, "Node financials", level=2)
    fin = pd.DataFrame(
        {
            "Metric": ["Revenue (USD bn)", "EBITDA (USD bn)", "EBITDA margin (%)"],
            "FY2022": [
                safe_num(node.get("segment_fy22_revenue_usd_bn")),
                safe_num(node.get("segment_fy22_ebitda_usd_bn")),
                safe_num(node.get("segment_fy22_ebitda_margin_pct")),
            ],
            "FY2023": [
                safe_num(node.get("segment_fy23_revenue_usd_bn")),
                safe_num(node.get("segment_fy23_ebitda_usd_bn")),
                safe_num(node.get("segment_fy23_ebitda_margin_pct")),
            ],
            "FY2024": [
                safe_num(node.get("segment_fy24_revenue_usd_bn")),
                safe_num(node.get("segment_fy24_ebitda_usd_bn")),
                safe_num(node.get("segment_fy24_ebitda_margin_pct")),
            ],
            "FY2025": [
                safe_num(node.get("segment_fy25_revenue_usd_bn")),
                safe_num(node.get("segment_fy25_ebitda_usd_bn")),
                safe_num(node.get("segment_fy25_ebitda_margin_pct")),
            ],
        }
    )
    # string formatting
    fin_fmt = fin.copy()
        fin_fmt[c] = [
            _fmt_num(value, 1 if "margin" in str(metric).lower() else 3)
            for metric, value in zip(fin["Metric"], fin[c])
        ]

    _doc_add_table(doc, fin_fmt, title="Financial summary")

    fin_comment = str(node.get("financial_commentary", "") or "").strip()
    if fin_comment:
        _doc_add_heading(doc, "Commentary", level=3)
        _doc_add_para(doc, fin_comment)

    # ----------------
    # Players & proxies (node filtered, FY25 snapshot)
    # ----------------
    _doc_add_heading(doc, "Players and proxies", level=2)
    pid = node_id
    if pid and "node_id" in players_all.columns:
        pl = players_all[players_all["node_id"].astype(str) == pid].copy()
    else:
        pl = pd.DataFrame()

    if not pl.empty:
        cols = []
        for c in ["player_name", "display_name", "player", "company"]:
            if c in pl.columns:
                cols.append(c)
                break
        # FY25 player metrics (if present)
        for c in ["fy25_revenue_usd_bn", "fy25_ebitda_usd_bn", "fy25_ebitda_margin_pct"]:
            if c in pl.columns:
                cols.append(c)
        # fallback if your player columns are prefixed differently
        if len(cols) == 1:
            # keep a reasonable view
            cols = pl.columns[:8].tolist()

        pl_view = pl[cols].head(15).copy()
        if "fy25_revenue_usd_bn" in pl_view.columns:
            pl_view["fy25_revenue_usd_bn"] = pl_view["fy25_revenue_usd_bn"].apply(lambda x: _fmt_num(x, 3))
        if "fy25_ebitda_usd_bn" in pl_view.columns:
            pl_view["fy25_ebitda_usd_bn"] = pl_view["fy25_ebitda_usd_bn"].apply(lambda x: _fmt_num(x, 3))
        if "fy25_ebitda_margin_pct" in pl_view.columns:
            pl_view["fy25_ebitda_margin_pct"] = pl_view["fy25_ebitda_margin_pct"].apply(lambda x: _fmt_num(x, 1))

        _doc_add_table(doc, pl_view, title="Players (top rows)")

        pl_comment = str(node.get("player_commentary", "") or "").strip()
        if pl_comment:
            _doc_add_heading(doc, "Commentary", level=3)
            _doc_add_para(doc, pl_comment)
    else:
        _doc_add_para(doc, "No player rows found for this node (or Players sheet missing node_id).")

    if pid and "node_id" in proxies_all.columns:
        pr = proxies_all[proxies_all["node_id"].astype(str) == pid].copy()
    else:
        pr = pd.DataFrame()
    if not pr.empty:
        pr_view = pr.copy()
        _doc_add_table(doc, pr_view.head(20), title="Proxies (top rows)")
    else:
        _doc_add_para(doc, "No proxy rows found for this node (or Proxies sheet missing node_id).")

    # ----------------
    # Evidence mapping (humanised)
    # ----------------
    _doc_add_heading(doc, "Evidence and audit trail", level=2)
    ev = pd.DataFrame()
    if isinstance(evidence_all, pd.DataFrame) and not evidence_all.empty and "node_id" in evidence_all.columns:
        ev = evidence_all[evidence_all["node_id"].astype(str) == pid].copy()

    if not ev.empty:
        # Choose a sane set of columns if present
        keep = []
        for c in ["evidence_id", "source_name", "source_title", "publisher", "url", "date", "snippet", "excerpt"]:
            if c in ev.columns:
                keep.append(c)
        if not keep:
            keep = ev.columns[:8].tolist()
        _doc_add_table(doc, ev[keep].head(25), title="Evidence (top rows)")
    else:
        _doc_add_para(doc, "No evidence rows found for this node.")

    em = pd.DataFrame()
    if isinstance(evidence_map_all, pd.DataFrame) and not evidence_map_all.empty:
        if "node_id" in evidence_map_all.columns:
            em = evidence_map_all[evidence_map_all["node_id"].astype(str) == pid].copy()
        else:
            # some versions store node reference in 'entity_id' / 'target_id'
            for alt in ["entity_id", "target_id", "ref_id"]:
                if alt in evidence_map_all.columns:
                    em = evidence_map_all[evidence_map_all[alt].astype(str) == pid].copy()
                    break

    if not em.empty:
        # Humanise support fields if present
        for col in ["field", "supports", "supported_field", "supported"]:
            if col in em.columns:
                em["Supports (human)"] = em[col].apply(lambda x: humanise_support_field(str(x)) if str(x).strip() else "")
                break
        keep = []
        for c in ["evidence_id", "Supports (human)", "field", "supports", "confidence", "note"]:
            if c in em.columns and c not in keep:
                keep.append(c)
        if not keep:
            keep = em.columns[:8].tolist()
        _doc_add_table(doc, em[keep].head(40), title="Evidence mapping (top rows)")
    else:
        _doc_add_para(doc, "No evidence mapping rows found for this node.")

    # Return bytes
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def render_custom_heatmaps(nodes: pd.DataFrame) -> None:
    """
    Old app approach:
      - user selects ROWS (taxonomy nodes) via filter + checkbox grid
      - user selects COLUMNS (metrics) via multiselect
      - click Generate to render heatmap
    """
    st.subheader("Custom heatmaps")
    if nodes is None or nodes.empty:
        st.info("No nodes loaded.")
        return
    if "path" not in nodes.columns:
        st.warning("Nodes sheet missing required column: 'path'")
        return

    # Derive hierarchy columns from `path` so we don't depend on explicit
    # 'Hierarchy - Level x' columns in the workbook.
    nodes, hier_cols = _with_path_hierarchy_from_df(nodes)

    if not hier_cols:
        st.warning("No hierarchy could be derived from 'path' for tree selection.")
        return

    # Only show rows that actually have a path
    nodes = nodes.copy()
    nodes["path"] = nodes["path"].astype(str)
    paths = [p for p in nodes["path"].dropna().unique().tolist() if p.strip()]
    paths.sort()

    controls, viz = st.columns([0.85, 2.15], gap="large")

    with controls:
        st.markdown("### Controls")

        # Make the button obvious and near the top (people kept missing it)
        generate = st.button("Generate heatmap", type="primary", use_container_width=True)

        st.markdown("#### Metrics (columns)")
        metric_options = list(CUSTOM_HEATMAP_METRICS.keys())
        default_cols = ["Margin (FY25)", "Revenue Growth 22–25", "EBITDA Growth 22–25"]
        selected_metrics = st.multiselect(
            "Select metrics",
            options=metric_options,
            default=[m for m in default_cols if m in metric_options],
        )

        st.markdown("#### Rows (taxonomy nodes)")

        # Persist selection in session state (list of full 'path' strings)
        if "custom_hm_rows" not in st.session_state:
            st.session_state.custom_hm_rows = []

        # Simple v1-style behaviour:
        # - use Level 1–3 dropdowns only to FILTER the list of paths
        # - pick actual nodes via a single multiselect over full paths
        with st.expander("Filter & pick nodes", expanded=True):
            scope = nodes
            # Level 1–3 filters (optional)
            for i, label in enumerate(["Level 1", "Level 2", "Level 3"]):
                if i >= len(hier_cols):
                    break
                col = hier_cols[i]
                vals = sorted(scope[col].dropna().astype(str).unique().tolist())
                choice = st.selectbox(label, options=["(all)"] + vals, index=0, key=f"hm_lvl{i+1}")
                if choice != "(all)":
                    scope = scope[scope[col].astype(str) == choice]

            # Available paths after filtering
            filtered_paths = sorted(
                p for p in scope["path"].dropna().astype(str).unique().tolist() if p.strip()
            )

            st.session_state.custom_hm_rows = st.multiselect(
                "Pick nodes for rows",
                options=filtered_paths,
                default=[p for p in st.session_state.custom_hm_rows if p in filtered_paths],
            )

        selected_rows = st.session_state.custom_hm_rows
        st.caption(f"Selected rows: **{len(selected_rows)}**")

        st.markdown("#### Colour scale")
        robust = st.checkbox("Robust scale (clip 5th–95th percentile)", value=True)
        symmetric = st.checkbox("Symmetric around 0", value=True)

    with viz:
        st.markdown("### Heatmap")
        st.caption("Tip: choose metrics + rows on the left, then click **Generate heatmap**.")

    # Only build chart when user clicks Generate (or if already generated once and we want it to persist)
    if generate:
        if not selected_rows:
            st.warning("Select at least one row.")
            return
        if not selected_metrics:
            st.warning("Select at least one metric column.")
            return

        # Lookup node rows by path
        lookup = nodes.set_index("path", drop=False)

        # Filter to selected nodes
        selected_nodes = nodes[nodes["path"].isin(selected_rows)].copy()

        # Derive hierarchy levels from path (old-app behaviour)
        max_depth = 4
        for i in range(max_depth):
            selected_nodes[f"Level {i+1}"] = selected_nodes["path"].apply(
                lambda p: split_path_levels(p, max_depth)[i]
                if len(split_path_levels(p, max_depth)) > i
                else None
            )

        y_labels = selected_rows
        x_labels = selected_metrics

        z = []
        hover = []

        for p in y_labels:
            r = lookup.loc[p] if p in lookup.index else None
            if r is None:
                z.append([np.nan] * len(x_labels))
                hover.append([f"{p}<br>{m}: (no data)" for m in x_labels])
                continue
            row_vals = []
            row_hover = []
            for m in x_labels:
                spec = CUSTOM_HEATMAP_METRICS[m]
                v = _metric_value(r, spec)
                row_vals.append(v if v is not None else np.nan)
                if v is None:
                    row_hover.append(f"{p}<br><b>{m}</b>: (no data)")
                else:
                    row_hover.append(f"{p}<br><b>{m}</b>: {v:.2f}{spec.get('unit','')}")
            z.append(row_vals)
            hover.append(row_hover)

        # Scale defaults (avoid one outlier making everything red)
        flat = [float(v) for row in z for v in row if not (isinstance(v, float) and math.isnan(v))]
        if not flat:
            st.warning("No numeric values available for the selected rows/metrics.")
            return

        if robust:
            zmin = float(np.percentile(flat, 5))
            zmax = float(np.percentile(flat, 95))
        else:
            zmin, zmax = float(min(flat)), float(max(flat))

        if symmetric:
            M = max(abs(zmin), abs(zmax))
            zmin, zmax = -M, M

        # Put min/max controls ABOVE the heatmap, but in the viz column so it doesn't clutter controls
        with viz:
            c1, c2 = st.columns(2)
            with c1:
                zmin_ui = st.number_input("Colour min", value=float(zmin), key="chm_zmin")
            with c2:
                zmax_ui = st.number_input("Colour max", value=float(zmax), key="chm_zmax")
        if zmin_ui < zmax_ui:
            zmin, zmax = float(zmin_ui), float(zmax_ui)

        fig = go.Figure(
            data=go.Heatmap(
                z=z,
                x=x_labels,
                y=y_labels,
                text=hover,
                hoverinfo="text",
                colorscale=MUTED_RAG,
                zmin=zmin,
                zmax=zmax,
                hoverongaps=False,
                colorbar=dict(title="Value"),
            )
        )
        fig.update_layout(
            height=min(1200, 24 * len(y_labels) + 260),
            margin=dict(l=10, r=10, t=60, b=10),
            title="Custom heatmap (selected rows × selected metrics)",
            paper_bgcolor="#0f172a",
            plot_bgcolor="#0f172a",
            font=dict(color="#e5e7eb"),
            hoverlabel=dict(
                bgcolor="rgba(15, 23, 42, 0.95)",
                bordercolor="rgba(148, 163, 184, 0.35)",
                font=dict(color="rgba(226, 232, 240, 1.0)", size=13),
            ),
            legend=dict(
                bgcolor="rgba(15, 23, 42, 0.9)",
                bordercolor="rgba(148, 163, 184, 0.35)",
                font=dict(color="rgba(226, 232, 240, 1.0)"),
            ),
        )
        with viz:
            st.plotly_chart(fig, use_container_width=True, theme=None)

def render_total_heatmap(nodes: pd.DataFrame) -> None:
    """
    Old-app style 'Heatmap - all':
      - rows: leaf paths (level == max(level))
      - cols: hierarchy levels (H1..Hn)
      - cell value: the chosen metric value for the ancestor node at that level
    """
    st.subheader("Total heatmap")

    if nodes is None or nodes.empty or "path" not in nodes.columns:
        st.info("No Nodes/path data available.")
        return

    # Leaf rows only (prevents diagonal stepping; matches old app behaviour)
    df_nodes = nodes.copy()
    if "level" in df_nodes.columns:
        lvl = pd.to_numeric(df_nodes["level"], errors="coerce")
        if lvl.notna().any():
            df_nodes = df_nodes[lvl == int(lvl.max())].copy()

    # Controls (match v1 intent, but use the new required options)
    metric_choice = st.selectbox(
        "Metric",
        [
            "Margin (FY23)",
            "Margin (FY24)",
            "Margin (FY25)",
            "Revenue Growth FY22–FY25",
            "EBITDA Growth FY22–FY25",
        ],
        index=2,  # default = Margin (FY25)
    )

    # Resolve choice to either a direct column or a derived series
    metric_kind = "margin"
    fmt_kind = "pct"  # pct | usd
    value_col = None

    if metric_choice.startswith("Margin"):
        metric_kind = "margin"
        fmt_kind = "pct"
        if "FY23" in metric_choice:
            value_col = "segment_fy23_ebitda_margin_pct"
        elif "FY24" in metric_choice:
            value_col = "segment_fy24_ebitda_margin_pct"
        else:
            value_col = "segment_fy25_ebitda_margin_pct"

        if value_col not in nodes.columns:
            st.warning(f"Missing column in Nodes: `{value_col}`")
            return

        # Build lookup from Nodes.path -> value
        node_lookup = (
            nodes[["path", value_col]]
            .copy()
            .assign(path=lambda d: d["path"].fillna("").astype(str).str.strip())
        )
        metric_map = dict(zip(node_lookup["path"], node_lookup[value_col].apply(safe_num)))

    elif metric_choice.startswith("Revenue Growth"):
        metric_kind = "rev_growth"
        fmt_kind = "pct"
        c0 = "segment_fy22_revenue_usd_bn"
        c1 = "segment_fy25_revenue_usd_bn"
        missing = [c for c in (c0, c1) if c not in nodes.columns]
        if missing:
            st.warning(f"Missing column(s) in Nodes: {', '.join(f'`{m}`' for m in missing)}")
            return

        # Build lookup from Nodes.path -> derived growth %
        base = nodes[["path", c0, c1]].copy()
        base["path"] = base["path"].fillna("").astype(str).str.strip()
        base[c0] = base[c0].apply(safe_num)
        base[c1] = base[c1].apply(safe_num)
        base["__val__"] = base.apply(lambda r: _pct_growth(r[c0], r[c1]), axis=1)
        metric_map = dict(zip(base["path"], base["__val__"]))

    else:  # EBITDA Growth FY22–FY25
        metric_kind = "ebitda_growth"
        fmt_kind = "pct"
        c0 = "segment_fy22_ebitda_usd_bn"
        c1 = "segment_fy25_ebitda_usd_bn"
        missing = [c for c in (c0, c1) if c not in nodes.columns]
        if missing:
            st.warning(f"Missing column(s) in Nodes: {', '.join(f'`{m}`' for m in missing)}")
            return

        base = nodes[["path", c0, c1]].copy()
        base["path"] = base["path"].fillna("").astype(str).str.strip()
        base[c0] = base[c0].apply(safe_num)
        base[c1] = base[c1].apply(safe_num)
        base["__val__"] = base.apply(lambda r: _pct_growth(r[c0], r[c1]), axis=1)
        metric_map = dict(zip(base["path"], base["__val__"]))

    # Leaf rows only (robust): use max depth from `path` not `level`
    df_nodes = nodes.copy()
    depths = df_nodes["path"].fillna("").astype(str).map(lambda p: len(_path_parts(p)))
    max_depth = int(depths.max()) if len(depths) else 0
    df_leaf_base = df_nodes[depths == max_depth].copy()

    # Build hierarchy columns from leaf paths
    df_leaf, hcols = _with_path_hierarchy_from_df(df_leaf_base)
    if not hcols:
        st.info("No hierarchy could be derived from `path`.")
        return

    # Build matrix Z: each row is a leaf path; each column is a hierarchy level
    z = []
    row_labels = []
    for _, r in df_leaf.iterrows():
        segs = [str(r.get(h, "") or "").strip() for h in hcols]
        segs = [s for s in segs if s]
        if not segs:
            continue
        row_labels.append(" > ".join(segs))
        row_vals = []
        for i in range(len(segs)):
            partial = " > ".join(segs[: i + 1])
            row_vals.append(_lookup_with_fallback(metric_map, partial))
        # pad to full width
        if len(row_vals) < len(hcols):
            row_vals += [float("nan")] * (len(hcols) - len(row_vals))
        z.append(row_vals)

    if not z:
        st.info("No heatmap rows available.")
        return

    # Plotly accepts list-of-lists; avoid numpy entirely
    z = [[float(v) if v is not None else float("nan") for v in row] for row in z]
    col_labels = [f"Level {i+1}" for i in range(len(hcols))]

    # -----------------------------
    # Color scale controls (client-proof)
    # -----------------------------
    # Flatten values ignoring NaN
    flat = []
    for row in z:
        for v in row:
            if isinstance(v, float) and math.isnan(v):
                continue
            flat.append(float(v))

    if not flat:
        st.info("No numeric values available for heatmap.")
        return

    # Default range: robust percentiles to avoid one outlier turning everything red
    # (common approach used in dashboards)
    try:
        p5 = float(np.percentile(flat, 5))
        p95 = float(np.percentile(flat, 95))
    except Exception:
        # fallback if percentile calc fails
        p5 = min(flat)
        p95 = max(flat)

    default_min = p5
    default_max = p95

    # For growth metrics, a symmetric scale around 0 is usually more interpretable
    is_growth = ("Growth" in metric_choice)
    if is_growth:
        M = max(abs(default_min), abs(default_max))
        default_min, default_max = -M, M

    with st.expander("Colour scale", expanded=False):
        c1, c2, c3 = st.columns([1.2, 1.2, 1.0], gap="large")
        with c1:
            robust = st.checkbox("Robust scale (clip to 5th–95th percentile)", value=True)
        with c2:
            symmetric = st.checkbox("Symmetric around 0", value=is_growth, disabled=not is_growth)
        with c3:
            st.caption("Tip: use robust + symmetric for growth metrics.")

        # If user disables robust, use full min/max
        if not robust:
            default_min = float(min(flat))
            default_max = float(max(flat))
            if is_growth and symmetric:
                M = max(abs(default_min), abs(default_max))
                default_min, default_max = -M, M

        zmin = st.number_input("Colour min", value=float(default_min))
        zmax = st.number_input("Colour max", value=float(default_max))

        # enforce sane ordering
        if zmin >= zmax:
            st.warning("Colour min must be less than colour max. Using defaults.")
            zmin, zmax = float(default_min), float(default_max)

        if is_growth and symmetric:
            M = max(abs(zmin), abs(zmax))
            zmin, zmax = -M, M

    # Hover text (keep it simple and client-friendly)
    hover = []
    for i, rp in enumerate(row_labels):
        row_hover = []
        for j, cl in enumerate(col_labels):
            v = z[i][j]
            if (isinstance(v, float) and math.isnan(v)):
                row_hover.append(f"{rp}<br>{cl}: n/a")
            else:
                if fmt_kind == "pct":
                    row_hover.append(f"{rp}<br>{cl}: {v:.1f}%")
                else:
                    row_hover.append(f"{rp}<br>{cl}: {v:.2f} (USD bn)")
        hover.append(row_hover)

    fig = go.Figure(
        data=go.Heatmap(
            z=z,
            x=col_labels,
            y=row_labels,
            text=hover,
            hoverinfo="text",
            colorscale=[
                [0.0,  "#7f1d1d"],
                [0.25, "#9a3412"],
                [0.5,  "#a16207"],
                [0.75, "#166534"],
                [1.0,  "#14532d"],
            ],
            zmin=zmin,
            zmax=zmax,
            hoverongaps=False,
        )
    )
    fig.update_layout(
        paper_bgcolor="#0f172a",
        plot_bgcolor="#0f172a",
        font=dict(color="#e5e7eb"),
        hoverlabel=dict(
            bgcolor="rgba(15, 23, 42, 0.95)",  # dark slate, slightly transparent
            bordercolor="rgba(148, 163, 184, 0.35)",
            font=dict(
                color="rgba(226, 232, 240, 1.0)",  # light text
                size=13,
            ),
        ),
        legend=dict(
            bgcolor="rgba(15, 23, 42, 0.9)",
            bordercolor="rgba(148, 163, 184, 0.35)",
            font=dict(color="rgba(226, 232, 240, 1.0)"),
        ),
        height=min(1200, 26 * len(row_labels) + 200),
        margin=dict(l=10, r=10, t=40, b=10),
        xaxis=dict(title="Hierarchy level"),
        yaxis=dict(title="Leaf path"),
    )
    st.plotly_chart(fig, use_container_width=True, theme=None)

    # Download as CSV (matrix with row label + columns)
    out = pd.DataFrame(z, columns=col_labels)
    out.insert(0, "Leaf path", row_labels)
    st.download_button(
        "Download heatmap matrix (CSV)",
        data=out.to_csv(index=False).encode("utf-8"),
        file_name=f"total_heatmap_{metric_choice.lower().replace(' ', '_').replace('–','-')}.csv",
        mime="text/csv",
    )

def _pick_node_index(nodes: pd.DataFrame, filters: dict) -> int | None:
    """Pick a node index matching filters (best-effort)."""
    if nodes is None or nodes.empty:
        return None
    df = nodes
    for k, v in filters.items():
        if k in df.columns and v is not None and str(v).strip() != "":
            df = df[df[k].astype(str).str.strip() == str(v).strip()]
    if df.empty:
        return None
    return int(df.index[0])

def render_taxonomy_architecture(nodes: pd.DataFrame) -> None:
    """
    Full taxonomy 'map' view (like v1): wide columns, tree-like layout.
    Uses Nodes.path (split on '>') and renders a row-per-leaf grid.
    """
    st.header("Taxonomy architecture")
    st.caption("Taxonomy map (click a node to select)")

    # Compact styling for the taxonomy map buttons (denser, less padding)
    st.markdown(
        """
        <style>
          /* tighter columns gap handled by Streamlit, but we tighten button padding/height */
          div[data-testid="stButton"] > button {
            padding: 0.15rem 0.45rem !important;
            min-height: 1.55rem !important;
            line-height: 1.1 !important;
            font-size: 0.85rem !important;
            border-radius: 0.4rem !important;
          }
          /* reduce vertical whitespace between blocks */
          .block-container { padding-top: 1.0rem; }
        </style>
        """,
        unsafe_allow_html=True,
    )

    if nodes is None or nodes.empty or "path" not in nodes.columns:
        st.info("No Nodes/path data available.")
        return

    # IMPORTANT: build the *map rows* from LEAF nodes only (old app behaviour).
    # In v2, `level` is reliable: max(level) == leaf level.
    df_map = nodes
    if "level" in nodes.columns:
        try:
            max_level = int(pd.to_numeric(nodes["level"], errors="coerce").max())
            df_map = nodes[pd.to_numeric(nodes["level"], errors="coerce") == max_level].copy()
        except Exception:
            df_map = nodes.copy()

    df, hcols = _with_path_hierarchy(df_map)
    if not hcols:
        st.info("No taxonomy paths found to build hierarchy.")
        return

    # Choose depth to render: v1 was 4 levels, but we support whatever exists.
    # Most client-friendly is 4; if more, we still render all (just more columns).
    depth = len(hcols)

    # Build leaf rows: we treat the deepest non-empty level per path as the leaf.
    # Row key = full path.
    df["_full_path"] = df["path"].fillna("").astype(str).str.strip()

    # Order LEAF rows by hierarchy to keep stable visual grouping
    leaf_rows = (
        df[hcols + ["_full_path"]]
        .fillna("")
        .astype(str)
        .drop_duplicates(subset=["_full_path"])
        .sort_values(by=hcols)
        .reset_index(drop=True)
    )

    # Track previous row values so we can suppress repeated ancestors (old-app style)
    prev = [""] * depth

    # Layout: slightly wider early levels
    weights = [1.2] + [1.2] + [1.2] + [1.4]
    if depth > 4:
        weights = weights + [1.2] * (depth - 4)

    # Render rows (compact matrix: no stepping/blank spacer rows)
    for ridx, r in leaf_rows.iterrows():
        row_cols = st.columns(weights[:depth], gap="small")

        for i, colname in enumerate(hcols):
            val = str(r.get(colname, "") or "").strip()

            # Build the path up to this level for selection
            # We reconstruct it from the row values up to i
            segs = [str(r.get(hcols[j], "") or "").strip() for j in range(i + 1)]
            segs = [s for s in segs if s]
            partial_path = " > ".join(segs)
            idx = _node_index_for_path(nodes, partial_path)

            with row_cols[i]:
                # Old-app layout:
                # - show the node only when it changes vs previous row at that column
                # - otherwise render blank to avoid repeating ancestors
                if not val:
                    st.markdown("<div style='height: 1.55rem;'></div>", unsafe_allow_html=True)
                    continue

                if val == prev[i]:
                    # blank cell (no repetition)
                    st.markdown("<div style='height: 1.55rem;'></div>", unsafe_allow_html=True)
                    continue

                # Show the node as a button
                if st.button(val, key=f"taxmap_{ridx}_{i}_{partial_path}"):
                    if idx is not None:
                        st.session_state["selected_idx"] = int(idx)
                        st.rerun()

        # update prev row tracker
        prev = [str(r.get(hcols[j], "") or "").strip() for j in range(depth)]


def _to_bullets(text: str) -> str:
    """
    Turn dense text into bullets.
    - If the text already contains newline-delimited bullets, keep them.
    - Otherwise split into sentences and bullet them.
    Returns HTML (escaped) inside <ul>.
    """
    if not text:
        return ""

    raw = str(text).strip()
    if not raw:
        return ""

    # If user/agent already gave newline-ish bullets, respect them
    lines = [ln.strip() for ln in re.split(r"\r?\n+", raw) if ln.strip()]
    looks_like_bullets = sum(1 for ln in lines if ln.startswith(("-", "•", "*"))) >= 2

    items: list[str] = []
    if looks_like_bullets:
        for ln in lines:
            ln2 = ln.lstrip("-•*").strip()
            if ln2:
                items.append(ln2)
    else:
        # Sentence split (good enough for now)
        sents = re.split(r"(?<=[.!?])\s+", raw)
        sents = [s.strip() for s in sents if s and s.strip()]
        # If it’s just one long sentence, keep as a single paragraph bullet
        items = sents if len(sents) > 1 else [raw]

    lis = "".join(f"<li>{html.escape(i)}</li>" for i in items)
    return f"<ul style='margin: 0.25rem 0 0.25rem 1.1rem;'>{lis}</ul>"

def _parse_evidence_snippet(snippet: str) -> tuple[str, str, str]:
    """
    Deterministic parser for the evidence memo format.
    We standardise around these labels:
      - 'Source reference:'
      - 'Quote:'
    Returns (source_reference, quote, remainder).
    If labels aren't present, quote falls back to the full snippet.
    """
    if not snippet:
        return ("", "", "")

    s = str(snippet).strip()
    if not s:
        return ("", "", "")

    # Normalise newlines
    s = re.sub(r"\r\n?", "\n", s)

    # Try to extract "Source reference:" and "Quote:" blocks
    # We keep it intentionally simple and tolerant to whitespace.
    src = ""
    quote = ""
    remainder = s

    m_src = re.search(r"(?is)\bSource\s+reference\s*:\s*(.*?)(?:\n\s*\bQuote\s*:|$)", s)
    if m_src:
        src = m_src.group(1).strip()

    m_q = re.search(r"(?is)\bQuote\s*:\s*(.*)$", s)
    if m_q:
        quote = m_q.group(1).strip()

    # If we extracted anything, compute a clean remainder (optional)
    if m_src or m_q:
        remainder = s
        if m_src:
            remainder = re.sub(r"(?is)\bSource\s+reference\s*:\s*.*?(?=\n\s*\bQuote\s*:|$)", "", remainder).strip()
        if m_q:
            remainder = re.sub(r"(?is)\bQuote\s*:\s*.*$", "", remainder).strip()

    # Hard fallback: if no Quote extracted, treat entire snippet as the quote body.
    if not quote:
        quote = s

    return (src, quote, remainder)


def render_card(title: str, text: str) -> None:
    """
    Non-scrolling “card” with bullet formatting for readability.
    """
    if not text:
        return
    body_html = _to_bullets(text)
    if not body_html:
        return

    st.markdown(
        f"""
        <div style="
            border: 1px solid rgba(255,255,255,0.08);
            border-radius: 10px;
            padding: 14px 16px;
            margin: 10px 0 18px 0;
            background: rgba(255,255,255,0.03);
        ">
          <div style="font-weight: 650; margin-bottom: 8px;">{html.escape(title)}</div>
          <div style="opacity: 0.92; line-height: 1.45;">{body_html}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


@st.cache_data(show_spinner=True)
def load_all_from_path(path: str) -> dict[str, pd.DataFrame]:
    return load_workbook_path(path)

@st.cache_data(show_spinner=True)
def load_all_from_bytes(xlsx_bytes: bytes) -> dict[str, pd.DataFrame]:
    return load_workbook_bytes(xlsx_bytes)

def postprocess(wb: dict[str, pd.DataFrame]) -> dict[str, pd.DataFrame]:
    nodes = wb.get("Nodes", pd.DataFrame()).copy()
    players = wb.get("Players", pd.DataFrame()).copy()
    proxies = wb.get("Proxies", pd.DataFrame()).copy()

    for col in ["node_id", "path", "display_name"]:
        if col not in nodes.columns:
            raise ValueError(f"Nodes sheet missing required column: {col}")

    nodes = derive_hierarchy(nodes)
    players = clean_players(players)
    proxies = clean_proxies(proxies)

    wb_out = dict(wb)
    wb_out["Nodes"] = nodes
    wb_out["Players"] = players
    wb_out["Proxies"] = proxies

    return wb_out

def _get_sheet(wb: dict[str, pd.DataFrame], names: list[str]) -> pd.DataFrame:
    """
    Return the first matching sheet as a DataFrame.
    IMPORTANT: don't use `df or other_df` because pandas DataFrames cannot be
    evaluated as booleans (raises ValueError: ambiguous truth value).
    """
    for name in names:
        df = wb.get(name)
        if isinstance(df, pd.DataFrame):
            return df
    return pd.DataFrame()

def _split_supported_by(x) -> list[str]:
    """
    Evidence_Map.supported_by is often stored as a string (e.g. "EVID:1, EVID:2")
    or sometimes a JSON-ish list. Keep it deterministic and forgiving.
    """
    if x is None:
        return []
    if isinstance(x, float) and pd.isna(x):
        return []
    if isinstance(x, (list, tuple, set)):
        return [str(i).strip() for i in x if str(i).strip()]
    s = str(x).strip()
    if not s:
        return []
    # Try JSON list first
    if s.startswith("[") and s.endswith("]"):
        try:
            import json
            arr = json.loads(s)
            if isinstance(arr, list):
                return [str(i).strip() for i in arr if str(i).strip()]
        except Exception:
            pass
    # Fallback: comma/semicolon separated
    parts = re.split(r"[;,]\s*|\s{2,}", s)
    return [p.strip() for p in parts if p.strip()]

def humanise_support_field(field: str) -> str:
    """
    Translate internal support fields into client-readable claims.
    Examples:
      fy25_revenue_usd_bn -> FY2025 Node Revenue
      player:Textron_Aviation:fy24_revenue_usd_bn -> FY2024 Textron Aviation Revenue
    """
    if not field:
        return field

    s = field.strip()

    # Player-level: player:Name:fyXX_metric
    if s.startswith("player:"):
        try:
            _, player, metric = s.split(":", 2)
            player = player.replace("_", " ")
            return f"{_humanise_fy_metric(metric)} ({player})"
        except Exception:
            return s

    # Node-level metric
    return _humanise_fy_metric(s)

def _humanise_fy_metric(metric: str) -> str:
    m = metric.lower()

    # Extract FY
    fy_match = re.search(r"fy(\d{2})", m)
    fy = f"FY20{fy_match.group(1)}" if fy_match else ""

    if "revenue" in m:
        label = "Revenue"
    elif "ebitda_margin" in m or "margin" in m:
        label = "EBITDA Margin"
    elif "ebitda" in m:
        label = "EBITDA"
    else:
        return metric

    scope = "Node"
    return f"{fy} {scope} {label}".strip()

def humanise_column_name(col: str) -> str:
    """
    Turn raw dataframe column names into client-friendly headers.
    Examples:
      player_fy25_revenue_usd_bn -> FY2025 Revenue (USD bn)
      proxy_fy24_ebitda_margin_pct -> FY2024 EBITDA Margin (%)
      confidence_score -> Confidence
      attribution_basis -> Attribution basis
    """
    if not col:
        return col
    s = str(col).strip()

    # Known non-metric columns
    static = {
        "rank": "Rank",
        "name": "Name",
        "country": "Country",
        "type": "Type",
        "proxy_reason": "Proxy rationale",
        "confidence_score": "Confidence",
        "attribution_basis": "Attribution basis",
        "node_id": "Node ID",
        "player_id": "Player ID",
        "proxy_id": "Proxy ID",
    }
    if s in static:
        return static[s]

    # Strip leading entity prefixes for readability in tables
    # player_fy25_revenue_usd_bn -> fy25_revenue_usd_bn
    # proxy_fy25_ebitda_usd_bn -> fy25_ebitda_usd_bn
    s2 = re.sub(r"^(player|proxy)_", "", s)

    # FY metric patterns
    m = re.match(r"^fy(\d{2})_(.+)$", s2.lower())
    if m:
        yy = m.group(1)
        rest = m.group(2)
        fy = f"FY20{yy}"

        if "revenue" in rest:
            return f"{fy} Revenue (USD bn)"
        if "ebitda_margin" in rest or rest.endswith("margin_pct") or "margin" in rest:
            return f"{fy} EBITDA Margin (%)"
        if "ebitda" in rest:
            return f"{fy} EBITDA (USD bn)"

    # Fallback: prettify snake_case
    return s.replace("_", " ").strip().title()

def render_node_header(node: pd.Series) -> None:
    """Consistent header at the top of every node-level tab."""
    name = str(node.get("display_name", "") or "").strip()
    path = str(node.get("path", "") or "").strip()
    if name:
        st.header(name)
    if path:
        st.caption(path)

def main() -> None:
    st.sidebar.header("Dataset")

    DEFAULT_PATH = "data/a_d_database.xlsx"

    uploaded = st.sidebar.file_uploader(
        "Optional: upload a different workbook",
        type=["xlsx"],
    )

    if uploaded is not None:
        wb = load_all_from_bytes(uploaded.read())
        label = f"Uploaded: {uploaded.name}"
    else:
        wb = load_all_from_path(DEFAULT_PATH)
        label = f"Default dataset: {DEFAULT_PATH}"

    wb = postprocess(wb)
    nodes = wb["Nodes"]
    players_all = wb["Players"]
    proxies_all = wb["Proxies"]
    evidence_all = _get_sheet(wb, ["Evidence", "EVIDENCE", "evidence"])
    evidence_map_all = _get_sheet(wb, ["Evidence_Map", "EVIDENCE_MAP", "evidence_map", "Evidence map"])

    st.title("A&D Market Explorer (v2)")
    st.caption(label)

    # ----------------------------
    # Taxonomy selector
    # ----------------------------
    st.sidebar.header("Taxonomy")
    q = st.sidebar.text_input("Search node", value="")

    nodes_view = nodes
    if q.strip():
        nodes_view = nodes[
            nodes["display_name"].astype(str).str.contains(q, case=False, na=False)
        ]

    if nodes_view.empty:
        st.sidebar.warning("No nodes match your search.")
        st.stop()

    labels = (
        nodes_view["display_name"].astype(str)
        + "  ·  "
        + nodes_view["path"].astype(str)
    )
    options = list(nodes_view.index)

    # Allow other parts of the app (taxonomy architecture) to set the selection
    if "selected_idx" not in st.session_state:
        st.session_state["selected_idx"] = int(options[0])
    if st.session_state["selected_idx"] not in options:
        st.session_state["selected_idx"] = int(options[0])

    selected_idx = st.sidebar.selectbox(
        "Select node",
        options=options,
        index=options.index(st.session_state["selected_idx"]),
        format_func=lambda i: labels.loc[i],
    )
    st.session_state["selected_idx"] = int(selected_idx)

    node = nodes.loc[selected_idx]

    # ----------------------------
    # KPI helper
    # ----------------------------
    def kpis(fy: int = 25):
        # Spreadsheet uses two-digit FY columns: segment_fy15 ... segment_fy25
        fy2 = fy % 100 if fy > 100 else fy
        fy_label = 2000 + fy2  # FY25 -> 2025 for display

        rev = safe_num(node.get(f"segment_fy{fy2:02d}_revenue_usd_bn"))
        ebitda = safe_num(node.get(f"segment_fy{fy2:02d}_ebitda_usd_bn"))
        margin = safe_num(node.get(f"segment_fy{fy2:02d}_ebitda_margin_pct"))

        c1, c2, c3 = st.columns(3)
        c1.metric(
            f"FY{fy_label} Revenue (USD bn)",
            f"{rev:,.3g}" if pd.notna(rev) else "—",
        )
        c2.metric(
            f"FY{fy_label} EBITDA (USD bn)",
            f"{ebitda:,.3g}" if pd.notna(ebitda) else "—",
        )
        c3.metric(
            f"FY{fy_label} EBITDA Margin (%)",
            f"{margin:,.3g}" if pd.notna(margin) else "—",
        )

    top_node, top_tax = st.tabs(["Node-level analysis", "Overall taxonomy analysis"])

    with top_node:
        tab1, tab2, tab3, tab4 = st.tabs(
            ["Overview", "Node Financials", "Players & Proxies", "Evidence mapping"]
        )

        with tab1:
            render_node_header(node)
            kpis()

            # ----------------------------
            # Node Word report (v1)
            # ----------------------------
            r1, r2 = st.columns([1.0, 2.0], gap="large")
            with r1:
                if st.button("Generate Word report (.docx)", type="primary", use_container_width=True):
                    if Document is None:
                        st.error(
                            "python-docx import failed. Add `python-docx` to requirements.txt.\n"
                            f"Details: {DOCX_IMPORT_ERROR}"
                        )
                    else:
                        try:
                            docx_bytes = build_node_report_docx(
                                node=node,
                                nodes=nodes,
                                players_all=players_all,
                                proxies_all=proxies_all,
                                evidence_all=evidence_all,
                                evidence_map_all=evidence_map_all,
                            )
                            st.session_state["node_report_docx"] = docx_bytes
                        except Exception as e:
                            st.error(f"Report generation failed: {e}")
            with r2:
                b = st.session_state.get("node_report_docx")
                if isinstance(b, (bytes, bytearray)) and b:
                    safe_name = str(node.get("display_name", "node")).strip().replace(" ", "_")
                    st.download_button(
                        "Download report",
                        data=b,
                        file_name=f"{safe_name}_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )

            # Overview content (no scroll boxes)
            # NOTE: financial_commentary moved to Tab 2 in the next step
            desc = str(node.get("node_description", "") or node.get("description", "") or "").strip()
            scope = str(node.get("scope_context", "") or "").strip()
            method = str(node.get("methodology_summary", "") or "").strip()

            if desc:
                render_card("Node description", desc)
            if scope:
                render_card("Scope", scope)
            if method:
                render_card("Methodology", method)

        with tab2:
            render_node_header(node)
            kpis()

            YEARS = list(range(15, 26))  # fy15 → fy25
            year_labels = [2000 + y for y in YEARS]

            # Build three series directly from the Nodes wide columns
            rev_cols = [f"segment_fy{y}_revenue_usd_bn" for y in YEARS]
            ebitda_cols = [f"segment_fy{y}_ebitda_usd_bn" for y in YEARS]
            margin_cols = [f"segment_fy{y}_ebitda_margin_pct" for y in YEARS]

            rev = [safe_num(node.get(c)) for c in rev_cols]
            ebitda = [safe_num(node.get(c)) for c in ebitda_cols]
            margin = [safe_num(node.get(c)) for c in margin_cols]

            # Combo chart: clustered bars (Revenue, EBITDA) + margin line on secondary axis
            x = year_labels
            fig = make_subplots(specs=[[{"secondary_y": True}]])

            fig.add_trace(
                go.Bar(
                    name="Revenue (USD bn)",
                    x=x,
                    y=rev,
                    offsetgroup="rev",
                    marker_color="#7CC7FF",
                ),
                secondary_y=False,
            )

            fig.add_trace(
                go.Bar(
                    name="EBITDA (USD bn)",
                    x=x,
                    y=ebitda,
                    offsetgroup="ebitda",
                    marker_color="#2E7BEF",
                ),
                secondary_y=False,
            )

            fig.add_trace(
                go.Scatter(
                    name="EBITDA Margin (%)",
                    x=x,
                    y=margin,
                    mode="lines+markers",
                    line=dict(color="#FF6B6B", width=2),
                    marker=dict(color="#FF6B6B"),
                ),
                secondary_y=True,
            )

            fig.update_layout(
                barmode="group",
                paper_bgcolor="#0f172a",
                plot_bgcolor="#0f172a",
                font=dict(color="#e5e7eb"),
                title=dict(
                    text="Revenue, EBITDA and Margin",
                    x=0.0,
                    xanchor="left",
                    y=0.98,
                    yanchor="top",
                ),
                margin=dict(l=10, r=10, t=55, b=10),
                legend=dict(
                    orientation="h",
                    yanchor="bottom",
                    y=1.02,
                    xanchor="left",
                    x=0,
                    bgcolor="rgba(15, 23, 42, 0.9)",
                    bordercolor="rgba(148, 163, 184, 0.35)",
                    font=dict(color="rgba(226, 232, 240, 1.0)"),
                ),
                hoverlabel=dict(
                    bgcolor="rgba(15, 23, 42, 0.95)",
                    bordercolor="rgba(148, 163, 184, 0.35)",
                    font=dict(color="rgba(226, 232, 240, 1.0)", size=13),
                ),
            )

            fig.update_xaxes(title_text="Fiscal Year", showgrid=False)

            # Kill horizontal gridlines (and all gridlines) - cleaner for clients
            fig.update_yaxes(title_text="USD bn", secondary_y=False, showgrid=False, zeroline=False)
            fig.update_yaxes(title_text="Margin (%)", secondary_y=True, showgrid=False, zeroline=False)

            st.plotly_chart(fig, use_container_width=True, theme=None)

            # FY table underneath (all three metrics)
            st.markdown("### FY15–FY25 table")

            # Transposed layout: years as columns (matches left→right chart flow)
            year_cols = [f"FY{y}" for y in year_labels]  # FY2015 ... FY2025
            table_t = pd.DataFrame(
                {
                    "Metric": ["Revenue (USD bn)", "EBITDA (USD bn)", "EBITDA Margin (%)"],
                    **{year_cols[i]: [rev[i], ebitda[i], margin[i]] for i in range(len(year_cols))},
                }
            )

            # Light formatting for readability
            def _fmt_fy_cell(metric: str, value) -> str:
                try:
                    if pd.isna(value):
                        return ""
                    decimals = 1 if "Margin" in str(metric) else 3
                    return f"{float(value):,.{decimals}f}"
                except Exception:
                    return ""

            table_fmt = table_t.copy()
            
            for i, y in enumerate(year_cols):
            table_fmt[y] = [
                    _fmt_fy_cell(metric, value)
                    for metric, value in zip(table_t["Metric"], table_t[y])
                ]

            st.dataframe(table_fmt, use_container_width=True, hide_index=True)

            # Financial commentary (below table)
            fin = str(node.get("financial_commentary", "") or "").strip()
            if fin:
                render_card("Financial commentary", fin)

        with tab3:
            render_node_header(node)
            st.markdown("### Players & Proxies")

            pid = str(node.get("node_id"))
            players = (
                players_all[players_all["node_id"].astype(str) == pid].copy()
                if (not players_all.empty and "node_id" in players_all.columns)
                else pd.DataFrame()
            )
            proxies = (
                proxies_all[proxies_all["node_id"].astype(str) == pid].copy()
                if (not proxies_all.empty and "node_id" in proxies_all.columns)
                else pd.DataFrame()
            )

            # ----------------------------
            # 1) Aggregated view (v1-style)
            # ----------------------------
            fy_pick = st.radio("Fiscal year", ["FY23", "FY24", "FY25"], index=2, horizontal=True)
            fy2 = int(fy_pick.replace("FY", ""))  # 23/24/25

            rev_col = f"player_fy{fy2}_revenue_usd_bn"
            ebitda_col = f"player_fy{fy2}_ebitda_usd_bn"
            mgn_col = f"player_fy{fy2}_ebitda_margin_pct"

            if players.empty:
                st.warning("No players found for this node.")
            else:
                p = players.copy()
                # Ensure numeric
                p[rev_col] = p[rev_col].map(safe_num)
                p[ebitda_col] = p[ebitda_col].map(safe_num)
                p[mgn_col] = p[mgn_col].map(safe_num)

                # Always sort by revenue descending for the selected fiscal year
                p = p.sort_values(rev_col, ascending=False)

                # Limit to keep charts readable (still "all players", but practical)
                p_chart = p.head(25).copy()
                # Sort once by revenue and keep this order for all three charts
                # Reverse the order so largest revenue appears at TOP of horizontal bar chart
                name_order = p_chart["name"].tolist()[::-1]

                c1, c2, c3 = st.columns(3)
                with c1:
                    fig1 = px.bar(
                        p_chart,
                        x=rev_col,
                        y="name",
                        orientation="h",
                        title=f"Revenue (USD bn) — {fy_pick}",
                        hover_data=["country", "type"],
                        color_discrete_sequence=["#7CC7FF"],
                    )
                    fig1.update_layout(
                        margin=dict(l=10, r=10, t=45, b=10),
                        paper_bgcolor="#0f172a",
                        plot_bgcolor="#0f172a",
                        font=dict(color="#e5e7eb"),
                        legend=dict(bgcolor="#0f172a", font=dict(color="#e5e7eb")),
                        hoverlabel=dict(bgcolor="#0f172a", font_color="#e5e7eb"),
                    )
                    fig1.update_yaxes(categoryorder="array", categoryarray=name_order)
                    st.plotly_chart(fig1, use_container_width=True, theme=None)

                with c2:
                    fig2 = px.bar(
                        p_chart,
                        x=ebitda_col,
                        y="name",
                        orientation="h",
                        title=f"EBITDA (USD bn) — {fy_pick}",
                        hover_data=["country", "type"],
                        color_discrete_sequence=["#2E7BEF"],
                    )
                    fig2.update_layout(
                        margin=dict(l=10, r=10, t=45, b=10),
                        paper_bgcolor="#0f172a",
                        plot_bgcolor="#0f172a",
                        font=dict(color="#e5e7eb"),
                        legend=dict(bgcolor="#0f172a", font=dict(color="#e5e7eb")),
                        hoverlabel=dict(bgcolor="#0f172a", font_color="#e5e7eb"),
                    )
                    fig2.update_yaxes(categoryorder="array", categoryarray=name_order)
                    st.plotly_chart(fig2, use_container_width=True, theme=None)

                with c3:
                    fig3 = px.bar(
                        p_chart,
                        x=mgn_col,
                        y="name",
                        orientation="h",
                        title=f"EBITDA Margin (%) — {fy_pick}",
                        hover_data=["country", "type"],
                        color_discrete_sequence=["#10B981"],
                    )
                    fig3.update_layout(
                        margin=dict(l=10, r=10, t=45, b=10),
                        paper_bgcolor="#0f172a",
                        plot_bgcolor="#0f172a",
                        font=dict(color="#e5e7eb"),
                        legend=dict(bgcolor="#0f172a", font=dict(color="#e5e7eb")),
                        hoverlabel=dict(bgcolor="#0f172a", font_color="#e5e7eb"),
                    )
                    fig3.update_yaxes(categoryorder="array", categoryarray=name_order)
                    st.plotly_chart(fig3, use_container_width=True, theme=None)

                st.caption("Charts show top 25 for readability; tables below include full rows for the node.")

                # Optional: full table directly beneath aggregated charts
                st.markdown("### Players table")
                p_tbl = p[[
                    "rank", "name", "country", "type",
                    "player_fy23_revenue_usd_bn", "player_fy24_revenue_usd_bn", "player_fy25_revenue_usd_bn",
                    "player_fy23_ebitda_usd_bn", "player_fy24_ebitda_usd_bn", "player_fy25_ebitda_usd_bn",
                    "player_fy23_ebitda_margin_pct", "player_fy24_ebitda_margin_pct", "player_fy25_ebitda_margin_pct",
                    "confidence_score", "attribution_basis"
                ]].copy()
                p_tbl = p_tbl.rename(columns={c: humanise_column_name(c) for c in p_tbl.columns})
                st.dataframe(p_tbl, use_container_width=True, hide_index=True)

            # ----------------------------
            # 2) Commentary
            # ----------------------------
            player_comm = str(node.get("player_commentary", "") or "").strip()
            proxy_comm = str(node.get("proxy_commentary", "") or "").strip()
            if player_comm:
                render_card("Player commentary", player_comm)
            if proxy_comm:
                render_card("Proxy commentary", proxy_comm)

            # ----------------------------
            # 3) Top 10 combo charts (FY23–FY25)
            # ----------------------------
            if not players.empty:
                st.markdown("### Top 10 players — FY23–FY25")
                p10 = players.copy()
                if "rank" in p10.columns and p10["rank"].notna().any():
                    p10 = p10.sort_values("rank", ascending=True)
                else:
                    p10 = p10.sort_values("player_fy25_revenue_usd_bn", ascending=False)
                p10 = p10.head(10)

                years = [2023, 2024, 2025]
                cols = st.columns(2)
                for i, (_, r) in enumerate(p10.iterrows()):
                    rev = [safe_num(r.get(f"player_fy{y%100}_revenue_usd_bn")) for y in years]
                    ebt = [safe_num(r.get(f"player_fy{y%100}_ebitda_usd_bn")) for y in years]
                    mgn = [safe_num(r.get(f"player_fy{y%100}_ebitda_margin_pct")) for y in years]

                    fig = make_subplots(specs=[[{"secondary_y": True}]])
                    fig.add_trace(go.Bar(name="Revenue (USD bn)", x=years, y=rev, offsetgroup="rev"), secondary_y=False)
                    fig.add_trace(go.Bar(name="EBITDA (USD bn)", x=years, y=ebt, offsetgroup="ebt"), secondary_y=False)
                    fig.add_trace(go.Scatter(name="Margin (%)", x=years, y=mgn, mode="lines+markers"), secondary_y=True)
                    fig.update_layout(
                        barmode="group",
                        paper_bgcolor="#0f172a",
                        plot_bgcolor="#0f172a",
                        font=dict(color="#e5e7eb"),
                        title=dict(text=str(r.get("name", "")), x=0.0, xanchor="left"),
                        margin=dict(l=10, r=10, t=55, b=10),
                        legend=dict(
                            orientation="h",
                            yanchor="bottom",
                            y=1.02,
                            xanchor="left",
                            x=0,
                            bgcolor="rgba(15, 23, 42, 0.9)",
                            bordercolor="rgba(148, 163, 184, 0.35)",
                            font=dict(color="rgba(226, 232, 240, 1.0)"),
                        ),
                        hoverlabel=dict(
                            bgcolor="rgba(15, 23, 42, 0.95)",
                            bordercolor="rgba(148, 163, 184, 0.35)",
                            font=dict(color="rgba(226, 232, 240, 1.0)", size=13),
                        ),
                    )
                    fig.update_xaxes(title_text="Fiscal Year", showgrid=False)
                    fig.update_yaxes(title_text="USD bn", secondary_y=False, showgrid=False, zeroline=False)
                    fig.update_yaxes(title_text="Margin (%)", secondary_y=True, showgrid=False, zeroline=False)

                    with cols[i % 2]:
                        st.plotly_chart(fig, use_container_width=True, theme=None)

            # Proxies table (simple for now; we can add proxy charts next)
            st.markdown("### Proxies table")
            if proxies.empty:
                st.caption("No proxies for this node.")
            else:
                pr_tbl = proxies[[
                    "name", "country", "type", "proxy_reason",
                    "proxy_fy23_revenue_usd_bn", "proxy_fy24_revenue_usd_bn", "proxy_fy25_revenue_usd_bn",
                    "proxy_fy23_ebitda_usd_bn", "proxy_fy24_ebitda_usd_bn", "proxy_fy25_ebitda_usd_bn",
                    "proxy_fy23_ebitda_margin_pct", "proxy_fy24_ebitda_margin_pct", "proxy_fy25_ebitda_margin_pct",
                    "confidence_score"
                ]].copy()
                pr_tbl = pr_tbl.rename(columns={c: humanise_column_name(c) for c in pr_tbl.columns})
                st.dataframe(pr_tbl, use_container_width=True, hide_index=True)

        with tab4:
            render_node_header(node)
            st.markdown("### Supporting evidence")

            pid = str(node.get("node_id", "") or "")
            if evidence_all.empty:
                st.warning("Evidence sheet is empty or missing from the workbook.")
                st.stop()

            # Expect evidence rows to include node_id + evidence_id
            if "node_id" not in evidence_all.columns or "evidence_id" not in evidence_all.columns:
                st.warning("Evidence sheet missing required columns: node_id and/or evidence_id.")
                st.stop()

            ev_pick = evidence_all[evidence_all["node_id"].astype(str) == pid].copy()
            if ev_pick.empty:
                st.info("No evidence rows found for this node_id.")
                st.stop()

            # Build an evidence_id -> [fields supported] mapping from Evidence_Map
            supports_by_evid: dict[str, list[str]] = {}
            if not evidence_map_all.empty and "node_id" in evidence_map_all.columns:
                em = evidence_map_all[evidence_map_all["node_id"].astype(str) == pid].copy()
                # Require at least field + supported_by
                if not em.empty and "field" in em.columns and "supported_by" in em.columns:
                    rows = []
                    for _, r in em.iterrows():
                        field = str(r.get("field", "") or "").strip()
                        evids = _split_supported_by(r.get("supported_by"))
                        for e in evids:
                            rows.append((e.strip(), field))
                    if rows:
                        tmp = pd.DataFrame(rows, columns=["evidence_id", "field"])
                        tmp["evidence_id"] = tmp["evidence_id"].astype(str).str.strip()
                        tmp["field"] = tmp["field"].astype(str).str.strip()
                        for evid, g in tmp.groupby("evidence_id"):
                            fields = [f for f in g["field"].tolist() if f]
                            # de-dupe preserving order
                            seen = set()
                            out = []
                            for f in fields:
                                if f not in seen:
                                    out.append(f)
                                    seen.add(f)
                            supports_by_evid[evid] = out

            def _clean(x):
                if x is None:
                    return ""
                if isinstance(x, float) and pd.isna(x):
                    return ""
                return str(x).strip()

            # Render evidence in a "research memo" style (like your screenshot)
            for _, r in ev_pick.iterrows():
                evid = _clean(r.get("evidence_id") or r.get("id") or "")
                title = _clean(r.get("title") or r.get("source_title") or "")
                url = _clean(r.get("url") or r.get("source_url") or "")

                # Optional richer fields (if present in other projects / future schema)
                dataset = _clean(r.get("dataset") or r.get("dataset_name") or "")
                src_type = _clean(r.get("type") or r.get("source_type") or "")
                supports = _clean(r.get("supports") or r.get("supports_field") or "")
                strength = _clean(r.get("strength") or r.get("strength_score") or "")
                conf = _clean(r.get("confidence_score") or "")

                # Canonical: parse from snippet so it is consistent across all evidence
                snippet = _clean(r.get("snippet") or r.get("excerpt") or "")
                source_ref, quote, _ = _parse_evidence_snippet(snippet)

                # What this evidence supports (from Evidence_Map)
                supports_fields = supports_by_evid.get(evid, [])
                supports_html = ""
                if supports_fields:
                    # Translate internal fields to client-readable claims
                    translated = [humanise_support_field(f) for f in supports_fields]

                    show_n = 8
                    shown = translated[:show_n]
                    more = len(supports_fields) - len(shown)
                    bullet_lis = "".join(f"<li>{html.escape(f)}</li>" for f in shown)
                    suffix = f"<div style='opacity:0.7; margin-top:6px;'>+{more} more</div>" if more > 0 else ""
                    supports_html = f"<div style='margin-top:10px;'><span style='font-weight:650;'>Supports:</span><ul style='margin:6px 0 0 1.1rem;'>{bullet_lis}</ul>{suffix}</div>"

                # Title line: clickable if URL available
                if title and url:
                    header_html = f"<a href='{html.escape(url)}' target='_blank' style='text-decoration:none; color: inherit;'>{html.escape(title)}</a>"
                elif title:
                    header_html = html.escape(title)
                else:
                    header_html = html.escape(evid or "Evidence item")

                # Prefix with evidence ID like "E16 — …"
                if evid:
                    header_html = f"<span style='opacity:0.85'>{html.escape(evid)}</span> — {header_html}"

                meta_parts = []
                if dataset:
                    meta_parts.append(dataset)
                if src_type:
                    meta_parts.append(src_type)
                if strength:
                    meta_parts.append(f"strength={strength}")
                if conf:
                    meta_parts.append(f"confidence={conf}")
                if supports:
                    meta_parts.append(f"supports: {supports}")
                meta = " · ".join(meta_parts)

                # Build sections: Source reference + Quote
                sr_html = ""
                if source_ref:
                    sr_html = f"<div style='margin-top: 8px;'><span style='font-weight:650;'>Source reference:</span> {html.escape(source_ref)}</div>"

                quote_html = ""
                if quote:
                    quote_html = f"<div style='margin-top: 6px;'><span style='font-weight:650;'>Quote:</span> {html.escape(quote)}</div>"

                st.markdown(
                    f"""
                    <div style="
                        border: 1px solid rgba(255,255,255,0.08);
                        border-radius: 12px;
                        padding: 16px 18px;
                        margin: 14px 0 16px 0;
                        background: rgba(255,255,255,0.02);
                    ">
                      <div style="font-weight: 750; font-size: 1.05rem; margin-bottom: 6px;">
                        {header_html}
                      </div>
                      {'<div style="opacity:0.75; font-size: 0.92rem; margin-bottom: 10px;">' + html.escape(meta) + '</div>' if meta else ''}
                      {sr_html}
                      {quote_html}
                      {supports_html}
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

    with top_tax:
        st.subheader("Overall taxonomy analysis")
        tax_tab1, tax_tab2, tax_tab3 = st.tabs(
            ["Taxonomy architecture", "Custom heatmaps", "Total heatmap"]
        )
        with tax_tab1:
            render_taxonomy_architecture(nodes)
        with tax_tab2:
            render_custom_heatmaps(nodes)
        with tax_tab3:
            render_total_heatmap(nodes)


if __name__ == "__main__":
    main()
